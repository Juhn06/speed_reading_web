# from flask import Flask, render_template, request, jsonify
# import os
# from docx import Document
# import PyPDF2
#
# app = Flask(__name__)
# app.config['UPLOAD_FOLDER'] = 'uploads'
#
# # Đảm bảo folder uploads tồn tại
# os.makedirs('uploads', exist_ok=True)
#
#
# # Trang chủ - Upload file
# @app.route('/', endpoint='main')  # Thêm endpoint='main'
# def index():
#     return render_template('home/index.html')
#
# # Trang đọc
# @app.route('/reader')
# def reader():
#     return render_template('reading/reader.html')
#
#
# # Xử lý upload file
# @app.route('/upload', methods=['POST'])
# def upload_file():
#     if 'file' not in request.files:
#         return jsonify({'error': 'Không có file'}), 400
#
#     file = request.files['file']
#     if file.filename == '':
#         return jsonify({'error': 'Chưa chọn file'}), 400
#
#     # Lưu file
#     filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
#     file.save(filepath)
#
#     # Đọc nội dung file
#     text = read_file(filepath)
#
#     # Tách thành từng từ
#     words = text.split()
#
#     return jsonify({
#         'words': words,
#         'total': len(words),
#         'filename': file.filename
#     })
#
#
# # Hàm đọc file theo loại
# def read_file(filepath):
#     ext = filepath.split('.')[-1].lower()
#
#     if ext == 'txt':
#         with open(filepath, 'r', encoding='utf-8') as f:
#             return f.read()
#
#     elif ext == 'docx':
#         doc = Document(filepath)
#         text = '\n'.join([para.text for para in doc.paragraphs])
#         return text
#
#     elif ext == 'pdf':
#         with open(filepath, 'rb') as f:
#             pdf = PyPDF2.PdfReader(f)
#             text = ''
#             for page in pdf.pages:
#                 text += page.extract_text()
#         return text
#
#     else:
#         return "Định dạng file không hỗ trợ"
#
# if __name__ == '__main__':
#     import os
#     port = int(os.environ.get('PORT', 5000))
#     app.run(host='0.0.0.0', port=port, debug=False)
from flask import Flask, render_template, request, jsonify, redirect, url_for
import os
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-here'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)


# Mock current_user for templates (since we don't have Flask-Login yet)
@app.context_processor
def inject_user():
    class MockUser:
        is_authenticated = False

    return dict(current_user=MockUser())


# Routes
@app.route('/', endpoint='main')
def index():
    return render_template('home/index.html')


@app.route('/reader')
def reader():
    return render_template('reading/reader.html')


@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'Không có file'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'error': 'Chưa chọn file'}), 400

    if file:
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        # Extract text based on file type
        text = extract_text(filepath)

        if text:
            words = text.split()
            return jsonify({
                'success': True,
                'filename': filename,
                'total_words': len(words),
                'words': words
            })
        else:
            return jsonify({'error': 'Không thể đọc file'}), 400


def extract_text(filepath):
    """Extract text from PDF or DOCX files"""
    text = ""

    if filepath.endswith('.pdf'):
        try:
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text()
        except Exception as e:
            print(f"Error reading PDF: {e}")
            return None

    elif filepath.endswith('.docx'):
        try:
            doc = Document(filepath)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            print(f"Error reading DOCX: {e}")
            return None

    elif filepath.endswith('.txt'):
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                text = file.read()
        except Exception as e:
            print(f"Error reading TXT: {e}")
            return None

    return text


# Mock routes for navbar links (to prevent errors)
@app.route('/auth/register', endpoint='auth.register')
def auth_register():
    return "Register page - Coming soon"


@app.route('/about', endpoint='main.about')
def about():
    return "About page - Coming soon"

@app.route('/auth/login', endpoint='auth.login')
def auth_login():
    return "Login page - Coming soon"


@app.route('/user/dashboard', endpoint='user.dashboard')
def user_dashboard():
    return "Dashboard - Coming soon"


if __name__ == '__main__':
    app.run(debug=True, host='127.0.0.1', port=5000)